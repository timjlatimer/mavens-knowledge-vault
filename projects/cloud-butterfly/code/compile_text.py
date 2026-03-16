from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_compilation_doc():
    doc = Document()
    
    # Title
    title = doc.add_heading('Cloud Butterfly & Solve Team: Text Compilation', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('This document contains all text-based protocols, definitions, and brainstorming materials for the Cloud Butterfly and Solve Team project, excluding Python code.')
    
    # --- SECTION 1: MASTER PROTOCOL ---
    doc.add_heading('1. Master Protocol (SOLVE Team Identity)', level=1)
    
    protocol_text = """
    Mission: We exist to solve the three "Gangsters" holding humanity back.
    1. Solve Capitalism (The Economy): "Cap The Gap" (Blue Seal) — Purpose with Profit.
    2. Solve Bureaucracy (The System): "Cut The Crap" (Done Seal) — Outcomes over Process.
    3. Solve Politics (The Leadership): "Track The Pact" (Promise Seal) — Accountability/Promises Kept.
    
    ENGINE A — THE LEARNING LOOP PROTOCOL
    Trigger: After every significant generation (draft, code, strategy, or analysis).
    
    The Instruction:
    You must Self-Evaluate your own work using this specific KPI prompt before asking the user for feedback.
    
    The Learning Loop KPI Check:
    1. Score (1-100): How well does this output match the user's brief and the Solve Team standards?
    2. The Gap: What is specifically missing or imperfect?
    3. The Fix: What is the one specific thing I will change in the next iteration to make it 10% better?
    4. Action: [Optional] If the score is below 80, automatically re-run the generation with the fix applied.
    """
    doc.add_paragraph(protocol_text)
    
    # --- SECTION 2: DESIGN & PERSONALITY ---
    doc.add_heading('2. Grace Design & Personality (Brainstorming)', level=1)
    
    design_text = """
    Design Movement: Neumorphic Futurism
    
    Core Principles:
    1. Soft Intelligence: Grace is AI, but feels organic. Soft shadows, rounded forms, and "breathing" animations.
    2. Tactile Interaction: Buttons and cards should feel pressable and responsive, mimicking physical controls.
    3. Clarity through Depth: Use layers and depth to organize information hierarchy, not just color.
    4. Personal Connection: The interface should feel like a dedicated workspace, not a generic dashboard.
    
    Color Philosophy ("Ethereal Dawn"):
    - Primary: Soft Periwinkle Blue (Intelligence, Calm)
    - Secondary: Warm Cream/Beige (Humanity, Approachability)
    - Accent: Golden Hour Yellow (Insight, Value)
    - Background: Off-white/Soft Grey (Clean slate)
    Reasoning: Grace is a "Cloud Butterfly" manager. The colors should evoke the sky, clouds, and the warmth of a personal assistant.
    
    Signature Elements:
    1. "Grace's Halo": A subtle, pulsing glow around the chat interface or active elements, representing her presence.
    2. "Butterfly Glass": Frosted glass effects (backdrop-filter) for overlays, mimicking the delicate nature of butterfly wings.
    3. "Fluid Connectors": Animated lines or particles that visually connect a "Butterfly" to a "Canvas" when an action is taken.
    """
    doc.add_paragraph(design_text)
    
    # --- SECTION 3: CANVAS DEFINITIONS ---
    doc.add_heading('3. Canvas Definitions (Text Content)', level=1)
    
    canvas_text = """
    PAGE 1: LEAN IMPACT CANVAS
    
    1. PROBLEM: Top 3 problems the user faces.
    2. CUSTOMER SEGMENTS: Target audience and early adopters.
    3. UNIQUE VALUE PROP: Single, clear, compelling message.
    4. SOLUTION: Top 3 features that solve the problems.
    5. UNFAIR ADVANTAGE: Something that cannot be easily copied or bought.
    6. REVENUE STREAMS: Revenue model, life time value, revenue, gross margin.
    7. COST STRUCTURE: Customer acquisition costs, distribution costs, hosting, people, etc.
    8. KEY METRICS: Key activities you measure.
    9. CHANNELS: Path to customers.
    10. IMPACT: Social and environmental impact.
    
    PAGE 2: STRATEGY EXECUTION (SOLVE)
    
    PURPOSE (BLUE SEAL): Why do we exist? (Beyond money)
    BHAG (DONE SEAL): Big Hairy Audacious Goal (10-25 years).
    CORE VALUES (PROMISE SEAL): Guiding principles.
    QUARTERLY PRIORITIES: Top 3-5 priorities for the next 90 days.
    KEY METRICS (KPIs): Critical numbers to watch.
    STRATEGIC MOVES (3Y / 1Y): Key milestones for the next 1-3 years.
    """
    doc.add_paragraph(canvas_text)
    
    doc.save('Cloud_Butterfly_Text_Compilation.docx')
    print("Document created successfully.")

if __name__ == "__main__":
    create_compilation_doc()
